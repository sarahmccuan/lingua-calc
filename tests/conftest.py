from __future__ import annotations

import io

import pytest
from docx import Document

from lingua_calc.models import ParsedToken, TokenFact

# Greek fixtures used across the suite. `Ὁ` vs `ὁ` is the case that motivated
# keeping a form breakdown: same lemma, same parse, different surface form.
HO = "ὁ"
HO_CAP = "Ὁ"
TOU = "τοῦ"
LOGOS = "λόγος"
LOGOU = "λόγου"
NOM = "nom. sg."
GEN = "gen. sg."


def fact(
    lemma: str,
    form: str,
    parse: str,
    chapter_index: int,
    position: int,
    *,
    type: str = "noun",
    filename: str = "a.docx",
) -> TokenFact:
    return TokenFact(
        type=type,
        lemma=lemma,
        form=form,
        parse=parse,
        filename=filename,
        chapter_index=chapter_index,
        chapter_id=f"{chapter_index + 1}-ch",
        chapter_title=f"Chapter {chapter_index + 1}",
        position=position,
    )


def facts_from(rows: list[tuple[str, str, str, int]], filename: str = "a.docx") -> list[TokenFact]:
    """Build facts from (lemma, form, parse, chapter_index) rows.

    Positions are assigned per chapter in the order given, matching how the
    pipeline numbers real provider output.
    """
    positions: dict[int, int] = {}
    out: list[TokenFact] = []
    for lemma, form, parse, chapter_index in rows:
        position = positions.get(chapter_index, 0)
        positions[chapter_index] = position + 1
        out.append(fact(lemma, form, parse, chapter_index, position, filename=filename))
    return out


class WordProvider:
    """Deterministic stand-in for Bedrock.

    Emits one token per whitespace-separated word. ``lemma_map`` lets a test
    point several surface forms at one lemma; anything unmapped lemmatizes to
    itself.
    """

    def __init__(self, lemma_map: dict[str, str] | None = None, parse: str = NOM) -> None:
        self.lemma_map = lemma_map or {}
        self.parse = parse
        self.calls: list[str] = []

    def analyze_chapter(self, text: str, chapter_title: str) -> list[ParsedToken]:
        self.calls.append(chapter_title)
        tokens: list[ParsedToken] = []
        for word in text.split():
            cleaned = word.strip(".,·;!?")
            if not cleaned:
                continue
            tokens.append(
                ParsedToken(
                    type="noun",
                    lemma=self.lemma_map.get(cleaned, cleaned),
                    form=cleaned,
                    parse=self.parse,
                )
            )
        return tokens


def make_docx(sections: list[tuple[str | None, str]]) -> bytes:
    """Build a .docx in memory. A non-None title becomes a Heading 1."""
    doc = Document()
    for title, body in sections:
        if title is not None:
            doc.add_heading(title, level=1)
        for para in body.split("\n\n"):
            doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def settings(tmp_path):
    """Settings pointed at a throwaway database."""
    from lingua_calc.config import Settings

    db_path = tmp_path / "test.sqlite3"
    settings = Settings(db_path=str(db_path), persist_runs=True, max_workers=2)

    # These fields are prefix-derived rather than alias-carrying, so a keyword
    # by field name binds. Give one a validation_alias and it silently stops
    # binding — which is how this fixture once handed back the default path and
    # the suite wrote to the real database. Assert rather than trust it.
    assert settings.db_path == str(db_path), "settings override was ignored"
    return settings
